import io
import os
import re

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Alignment
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from rest_framework import status
from rest_framework.response import Response
from rest_framework.views import APIView

from .models import GeneratedFile, OCRRecord, persist_generated_export

try:
    import arabic_reshaper
    from bidi.algorithm import get_display
except Exception:
    arabic_reshaper = None
    get_display = None


FONT_NAME = 'ArialUnicode'


def get_accessible_ocr_queryset(user):
    queryset = OCRRecord.objects.select_related('user')
    if user.is_staff or user.is_superuser:
        return queryset
    return queryset.filter(user=user)


def split_text_into_paragraphs(raw_text):
    text = (raw_text or '').strip()
    if not text:
        return []

    if '\n\n' in text:
        parts = text.split('\n\n')
    elif '\n' in text:
        parts = text.split('\n')
    else:
        return [{'index': 1, 'text': text}]

    paragraphs = []
    for part in parts:
        cleaned = part.strip()
        if cleaned:
            paragraphs.append({'index': len(paragraphs) + 1, 'text': cleaned})

    if not paragraphs:
        return []

    return paragraphs


def build_table_rows(raw_text):
    rows = []
    for line in (raw_text or '').splitlines():
        stripped_line = line.strip()
        if not stripped_line:
            continue

        cells = [cell.strip() for cell in re.split(r'\s{2,}', stripped_line) if cell and cell.strip()]
        if not cells:
            continue
        rows.append(cells)
    return rows


def to_pdf_text(value):
    text = value or ''
    try:
        if arabic_reshaper and get_display:
            return get_display(arabic_reshaper.reshape(text))
    except Exception:
        pass
    return text


def to_word_text(value):
    return value or ''


def set_rtl_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()

    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), 'on')
    p_pr.append(bidi)

    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    p_pr.append(jc)


class OCRExportPdfView(APIView):
    def get(self, request, pk, *args, **kwargs):
        ocr_record = get_object_or_404(get_accessible_ocr_queryset(request.user), pk=pk)

        if ocr_record.status != OCRRecord.STATUS_COMPLETED:
            return Response(
                {
                    'message': 'OCR record is not completed yet.',
                    'status': ocr_record.status,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        font_path = r'C:\Windows\Fonts\arial.ttf'
        if not os.path.exists(font_path):
            return Response(
                {'message': f'Arabic font not found at {font_path}.'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        try:
            if FONT_NAME not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(FONT_NAME, font_path))
        except Exception:
            return Response(
                {'message': 'Failed to load Arabic PDF font.'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        if not (arabic_reshaper and get_display):
            return Response(
                {'message': 'Arabic shaping libraries are missing. Install arabic-reshaper and python-bidi.'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=30,
            rightMargin=30,
            topMargin=30,
            bottomMargin=30,
        )
        story = []

        document_type = getattr(ocr_record, 'document_type', 'page') or 'page'

        if document_type == 'table':
            rows = build_table_rows(ocr_record.extracted_text)
            if not rows:
                cells = [to_pdf_text(ocr_record.extracted_text or '')]
                rows = [cells]

            max_columns = max(len(row) for row in rows)
            normalized_rows = []
            for row in rows:
                padded = [to_pdf_text(cell) for cell in row]
                if len(padded) < max_columns:
                    padded.extend([''] * (max_columns - len(padded)))
                normalized_rows.append(padded)

            table = Table(normalized_rows, repeatRows=1)
            table.setStyle(
                TableStyle([
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#EAEAEA')),
                    ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 0), (-1, -1), FONT_NAME),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F7F7F7')]),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ])
            )
            story.append(table)
        else:
            paragraphs = split_text_into_paragraphs(ocr_record.extracted_text)
            styles = getSampleStyleSheet()
            paragraph_style = ParagraphStyle(
                'RTLExportParagraph',
                parent=styles['Normal'],
                fontName=FONT_NAME,
                fontSize=12,
                leading=18,
                alignment=2,
                textColor=colors.black,
            )
            for paragraph in paragraphs:
                story.append(Paragraph(to_pdf_text(paragraph.get('text', '')), paragraph_style))
                story.append(Spacer(1, 10))

        doc.build(story)
        buffer.seek(0)
        pdf_bytes = buffer.getvalue()
        persist_generated_export(
            ocr_record,
            GeneratedFile.FileType.PDF,
            pdf_bytes,
            f'ocr_export_{ocr_record.id}.pdf',
        )

        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="ocr_export_{ocr_record.id}.pdf"'
        return response


class OCRExportDocxView(APIView):
    def get(self, request, pk, *args, **kwargs):
        ocr_record = get_object_or_404(get_accessible_ocr_queryset(request.user), pk=pk)

        if ocr_record.status != OCRRecord.STATUS_COMPLETED:
            return Response(
                {
                    'message': 'OCR record is not completed yet.',
                    'status': ocr_record.status,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        font_path = r'C:\Windows\Fonts\arial.ttf'
        if not os.path.exists(font_path):
            return Response(
                {'message': f'Arabic font not found at {font_path}.'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        if not (arabic_reshaper and get_display):
            return Response(
                {'message': 'Arabic shaping libraries are missing. Install arabic-reshaper and python-bidi.'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        doc = Document()
        default_style = doc.styles['Normal']
        default_style.font.name = 'Arial'
        default_style.font.size = Pt(12)
        default_style.rtl = True

        try:
            rpr = default_style.element.rPr
            if rpr is not None:
                rfonts = rpr.rFonts
                if rfonts is None:
                    rfonts = OxmlElement('w:rFonts')
                    rpr.append(rfonts)
                rfonts.set(qn('w:eastAsia'), 'Arial')
        except Exception:
            pass

        document_type = getattr(ocr_record, 'document_type', 'page') or 'page'

        if document_type == 'table':
            rows = build_table_rows(ocr_record.extracted_text)
            if not rows:
                rows = [[to_word_text(ocr_record.extracted_text or '')]]

            max_columns = max((len(row) for row in rows), default=1)
            normalized_rows = []
            for row in rows:
                cells = [to_word_text(cell) for cell in row]
                if len(cells) < max_columns:
                    cells.extend([''] * (max_columns - len(cells)))
                normalized_rows.append(cells)

            table = doc.add_table(rows=len(normalized_rows), cols=max_columns)
            table.style = 'Table Grid'
            for row_idx, row in enumerate(normalized_rows):
                for col_idx, cell_value in enumerate(row):
                    cell = table.cell(row_idx, col_idx)
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    if paragraph.runs:
                        paragraph.runs[0].text = cell_value
                    else:
                        paragraph.add_run(cell_value)
                    set_rtl_paragraph(paragraph)
        else:
            paragraphs = split_text_into_paragraphs(ocr_record.extracted_text)
            for paragraph_data in paragraphs:
                paragraph = doc.add_paragraph()
                paragraph_text = to_word_text(paragraph_data.get('text', ''))
                paragraph.add_run(paragraph_text)
                set_rtl_paragraph(paragraph)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.getvalue()
        persist_generated_export(
            ocr_record,
            GeneratedFile.FileType.WORD,
            docx_bytes,
            f'ocr_export_{ocr_record.id}.docx',
        )

        response = HttpResponse(
            docx_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = f'attachment; filename="ocr_export_{ocr_record.id}.docx"'
        return response


class OCRExportXlsxView(APIView):
    def get(self, request, pk, *args, **kwargs):
        ocr_record = get_object_or_404(get_accessible_ocr_queryset(request.user), pk=pk)

        if ocr_record.status != OCRRecord.STATUS_COMPLETED:
            return Response(
                {
                    'message': 'OCR record is not completed yet.',
                    'status': ocr_record.status,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        wb = Workbook()
        ws = wb.active
        ws.sheet_view.rightToLeft = True

        document_type = getattr(ocr_record, 'document_type', 'page') or 'page'

        if document_type == 'table':
            rows = build_table_rows(ocr_record.extracted_text)
            if not rows:
                rows = [[to_word_text(ocr_record.extracted_text or '')]]

            max_columns = max((len(row) for row in rows), default=1)
            normalized_rows = []
            for row in rows:
                cells = [to_word_text(cell) for cell in row]
                if len(cells) < max_columns:
                    cells.extend([''] * (max_columns - len(cells)))
                normalized_rows.append(cells)

            for row_idx, row in enumerate(normalized_rows, start=1):
                for col_idx, value in enumerate(row, start=1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.alignment = Alignment(horizontal='right')
        else:
            paragraphs = split_text_into_paragraphs(ocr_record.extracted_text)
            for row_idx, paragraph in enumerate(paragraphs, start=1):
                cell_value = to_word_text(paragraph.get('text', ''))
                cell = ws.cell(row=row_idx, column=1, value=cell_value)
                cell.alignment = Alignment(horizontal='right')

        for column_cells in ws.columns:
            ws.column_dimensions[column_cells[0].column_letter].width = 30

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        xlsx_bytes = buffer.getvalue()
        persist_generated_export(
            ocr_record,
            GeneratedFile.FileType.EXCEL,
            xlsx_bytes,
            f'ocr_export_{ocr_record.id}.xlsx',
        )

        response = HttpResponse(
            xlsx_bytes,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        response['Content-Disposition'] = f'attachment; filename="ocr_export_{ocr_record.id}.xlsx"'
        return response
